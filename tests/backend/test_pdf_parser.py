from pathlib import Path

from PIL import Image, ImageDraw
from docx import Document as WordDocument
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from backend.core.ingestion.pdf_parser import IntelliCreditPDFParser


def test_pdf_parser_detects_annual_report(tmp_path: Path):
    pdf_path = tmp_path / "annual_report.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=A4)
    c.drawString(40, 800, "Directors' Report")
    c.drawString(40, 780, "Standalone Financial Statements")
    c.drawString(40, 760, "Revenue from Operations 125000000")
    c.drawString(40, 740, "Profit After Tax 8400000")
    c.save()

    parsed = IntelliCreditPDFParser().parse(str(pdf_path))
    assert parsed.document_type.value in {"ANNUAL_REPORT", "FINANCIAL_STATEMENT"}
    assert parsed.extraction_confidence > 0.5


def test_parser_supports_jpeg_input(tmp_path: Path):
    image_path = tmp_path / "scan.jpg"
    image = Image.new("RGB", (1200, 800), color="white")
    draw = ImageDraw.Draw(image)
    draw.text((40, 40), "Balance Sheet\nRevenue from Operations 1000000", fill="black")
    image.save(image_path)

    parsed = IntelliCreditPDFParser().parse(str(image_path))
    assert parsed.company_name is not None
    assert 0.0 <= parsed.extraction_confidence <= 1.0


def test_parser_supports_docx_input(tmp_path: Path):
    docx_path = tmp_path / "annual_report.docx"
    doc = WordDocument()
    doc.add_paragraph("Directors' Report")
    doc.add_paragraph("Standalone Financial Statements")
    doc.add_paragraph("Revenue from Operations 125000000")
    doc.add_paragraph("Profit After Tax 8400000")
    doc.save(str(docx_path))

    parsed = IntelliCreditPDFParser().parse(str(docx_path))
    assert parsed.document_type.value in {"ANNUAL_REPORT", "FINANCIAL_STATEMENT"}
    assert parsed.extraction_confidence > 0.5
